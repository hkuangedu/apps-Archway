"""Generate the Archway agent planning & action plan as a .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "Archway_Agent_Planning.docx"

doc = Document()

# Base styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)


def h2(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)


def h3(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def para(text: str) -> None:
    doc.add_paragraph(text)


def bullets(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def mono_block(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(rows: list[list[str]], header: bool = True) -> None:
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.text = val
            if header and r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


# Cover
h1("Archway — Agent Planning & Action Plan")
p = doc.add_paragraph()
p.add_run("Student-Managed Fund | AI Workflow Sandbox").italic = True
p = doc.add_paragraph()
p.add_run("Last updated: 2026-05-28").italic = True
doc.add_paragraph()

# Executive Summary
h2("1. Executive Summary")
para(
    "Archway is a sandbox for the student-managed fund to learn how AI can improve "
    "research, monitoring, and reporting workflows. The v0 dashboard (watchlist, event calendar, "
    "rating tracker) is shipped. The next phase moves from passive dashboards to active agents — "
    "starting with a single earnings-briefing agent and growing into a multi-agent reporting system."
)
para(
    "Architectural decision: build runtime agents on Claude Managed Agents (launched April 2026) "
    "rather than hand-rolling orchestration. This reduces maintenance burden — critical for a "
    "student fund where contributors turn over every 1–4 years — and gives students exposure to "
    "current industry patterns instead of legacy plumbing."
)

# Current State
h2("2. Current State (Phase 0 — Shipped)")
bullets([
    "Python + Streamlit multi-page application, repo at github.com/hkuangedu/apps-Archway",
    "Watchlist page — sector-organized, CSV-backed, editable via st.data_editor",
    "Event Calendar page — upcoming earnings (Finnhub)",
    "Rating Tracker page — current sell-side recommendation distribution (Finnhub)",
    "Deployed on Streamlit Community Cloud with Finnhub free-tier API key",
])

# Stack
h2("3. Stack Decisions")
table([
    ["Layer", "Choice", "Why"],
    ["Building the system", "Claude Code", "Students pair with Claude to write the app"],
    ["Runtime agents", "Claude Managed Agents", "Hosted; built-in grader, sandbox, memory; less to maintain"],
    ["Daily fund workflow", "Claude Cowork (Excel/Word/PPT)", "Meets students where they already work"],
    ["Collaboration plan", "Claude Team (5–75 seats)", "Shared Projects, connectors, fund-wide billing"],
    ["App layer", "Python + Streamlit", "Idea → working tool in minutes; aligns with practitioner stacks"],
    ["Data (v0)", "Finnhub free tier", "Earnings calendar + recommendation trends"],
    ["Data (later)", "Fund's existing vendor (TBD)", "Pluggable via archway/data/ adapter pattern"],
    ["Hosting", "Streamlit Community Cloud", "Free; school infra is the fallback"],
])

# Agent Planning Map
h2("4. Agent Planning Map")
para(
    "Agents are grouped by role. Producers do the work; orchestrators route it; quality agents "
    "police it; formatters package it for the audience."
)

h3("Layer 1 — Producer Agents")
table([
    ["Agent", "Job", "Tools needed"],
    ["Earnings Briefing", "Per-ticker briefing on day of report", "get_calendar, get_actuals, get_press_release, get_news"],
    ["Sector Commentary", "Sector view for the quarterly report", "get_sector_returns, get_top_movers, get_macro_context"],
    ["Performance", "Fund-level returns, attribution, top contributors/detractors", "get_portfolio, get_returns, get_attribution"],
    ["Holdings", "Position-level commentary, concentration, changes vs prior period", "get_positions, get_changes"],
    ["Market", "Macro / market commentary for the period", "get_indices, get_rates, get_macro_news"],
])

h3("Layer 2 — Orchestration")
bullets([
    "Orchestrator Agent — reads the report template, dispatches sub-agents in parallel, assembles the draft",
    "In Managed Agents, this is the 'lead agent' pattern with specialist subagents on a shared filesystem",
])

h3("Layer 3 — Quality Agents")
bullets([
    "Verification Agent — re-reads each numeric claim against the cited source; loops back to producer if a number doesn't tie",
    "  Implemented as the Managed Agents rubric grader (built-in, no custom code)",
    "Governance Agent — enforces fund rules: must-cite-source, no investment-advice language, mandate alignment, MNPI scrubbing",
    "  Sits as a checkpoint between producer output and user/file",
])

h3("Layer 4 — Output")
bullets([
    "Formatter Agent — renders the final draft into Word / PowerPoint / PDF per template",
    "Uses python-docx / python-pptx; the LLM produces structured content, code produces the file",
])

h3("Flow (Reporting Example)")
para(
    "The reporting workflow chains agents from data → draft → quality → delivery. "
    "The orchestrator dispatches sector/data sub-agents in parallel; the writer assembles the draft; "
    "verification loops back to the writer if numbers don't tie; governance allows/rewrites/blocks "
    "the final language; the formatter produces the file."
)
mono_block(
    "                  ┌─────────────────┐\n"
    "                  │   Orchestrator  │  reads template, dispatches work\n"
    "                  └────────┬────────┘\n"
    "                           │\n"
    "         ┌─────────────────┼─────────────────┐\n"
    "         ▼                 ▼                 ▼\n"
    "   Performance Agent  Holdings Agent   Sector Agents (Tech, Health, ...)\n"
    "         │                 │                 │\n"
    "         └─────────────────┼─────────────────┘\n"
    "                           ▼\n"
    "                  ┌─────────────────┐\n"
    "                  │   Writer Agent  │  assembles draft from sections\n"
    "                  └────────┬────────┘\n"
    "                           ▼\n"
    "                  ┌─────────────────┐\n"
    "                  │ Verification    │  every number tied back to source\n"
    "                  │     Agent       │  → loops back to Writer if breaks\n"
    "                  └────────┬────────┘\n"
    "                           ▼\n"
    "                  ┌─────────────────┐\n"
    "                  │  Governance     │  language, mandate, disclosures\n"
    "                  │     Agent       │  → allow/rewrite/block\n"
    "                  └────────┬────────┘\n"
    "                           ▼\n"
    "                  ┌─────────────────┐\n"
    "                  │ Formatter Agent │  → Word / PPT / PDF\n"
    "                  └─────────────────┘"
)

# Action Plan
h2("5. Action Plan")

h3("Phase 1 — First Agent: Earnings Briefing  (Weeks 1–3)")
para("Goal: a single agent reads the watchlist + today's calendar and produces a 1-page briefing per reporting name.")
numbered([
    "Acquire Anthropic API access; add ANTHROPIC_API_KEY to Streamlit Cloud secrets",
    "Set up a Managed Agents project; configure base model (suggest Sonnet for cost)",
    "Implement Python tool functions: get_todays_earnings, get_actuals, get_press_release (SEC EDGAR), get_recent_news",
    "Write the agent system prompt: persona, output format (markdown briefing), tone (research, not advice)",
    "Add a Streamlit page 'Morning Briefing' with a 'Run today's briefings' button and a per-ticker output panel",
    "Add Governance v1: a single rule — every numeric claim must include a parenthetical source citation",
    "Test on 5–10 names from the current watchlist; iterate on the prompt",
])
para("Success criteria: clicking the button produces a cited, on-tone briefing for every name reporting today. No fabricated numbers.")

h3("Phase 2 — Verification + Polish  (Weeks 4–6)")
numbered([
    "Define the rubric: 'every numeric claim cites a source', 'tone is research not advice', 'no MNPI', 'within 350 words'",
    "Configure Managed Agents' grader with the rubric; enable revise-on-fail loop",
    "Expand Governance rules: mandate alignment, prohibited language list, source whitelist",
    "Add UI affordances: 'good briefing' / 'needs revision' feedback buttons (feed into prompt improvement)",
    "Document the prompts and rubric in the repo so the next student cohort can iterate",
])

h3("Phase 3 — Reporting System  (Weeks 7–12)")
numbered([
    "Inventory current process: locate an actual quarterly/monthly report template the fund uses today",
    "Decide output format (recommend PPT for board/donor meetings)",
    "Identify fund data source (custodian API, vendor, Excel) and build the data tools",
    "Build the orchestrator agent (lead agent in Managed Agents)",
    "Build sector sub-agents — one per sector team; each writes its own slide/section",
    "Wire shared Verification + Governance from Phase 2 over the reporting outputs",
    "Build the Formatter agent (python-pptx) to assemble the final deck",
    "Pilot for one reporting cycle in parallel with the existing manual process; compare",
])

h3("Phase 4 — Operational Maturity  (Ongoing)")
bullets([
    "Memory curation via Managed Agents' Dreaming — captures recurring patterns for reuse",
    "Webhooks → Slack/email when scheduled briefings or reports complete",
    "Self-hosted sandbox if fund data privacy requires staying inside school infra",
    "Skills library — package recurring patterns (e.g., 'guidance-change detector') as reusable Skills",
    "Cost monitoring + budget alerts; per-user rate limits",
])

# Open Questions
h2("6. Open Questions Before Phase 1")
numbered([
    "Where does fund-level position/return data live today (custodian, Bloomberg PORT, Excel)?",
    "Is there an existing report template the fund uses today? Can we get a copy?",
    "Who funds API usage — school budget, fund operating budget, advisor's existing account?",
    "Should Archway stay public on Streamlit Cloud, or move to a private app / school infra as agents handle more sensitive content?",
    "How are sector commentaries written today (Word, Slack, email) — what is the agent replacing?",
])

# Risks
h2("7. Risks & Mitigations")
table([
    ["Risk", "Likelihood", "Mitigation"],
    ["LLM hallucinates a number in a briefing", "High", "Verification agent (rubric grader) is non-negotiable; cite-the-source rule in Governance"],
    ["Cost runaway from open agent triggers", "Medium", "Access controls; per-user rate limits; private Streamlit app once agents go live"],
    ["Vendor lock-in to Anthropic", "Accepted", "Thin abstraction over the agent layer; tools stay provider-neutral"],
    ["Knowledge loss as students graduate", "High", "Managed Agents reduces hand-rolled code; document prompts and rubrics in /docs"],
    ["MNPI / compliance exposure", "Low (public data only at v0)", "Governance agent scrubs and logs; revisit before adding real fund data"],
])

# Educational Goals
h2("8. Educational Outcomes by Phase")
table([
    ["Phase", "What students learn"],
    ["Phase 0 (done)", "Web app stack, data plumbing, Streamlit, deploying to cloud"],
    ["Phase 1", "Agent loop, tool design, prompting, structured outputs"],
    ["Phase 2", "Evaluation rubrics, LLM-as-judge, self-critique loops"],
    ["Phase 3", "Multi-agent orchestration, separation of concerns, document generation"],
    ["Phase 4", "Production patterns: observability, scheduling, cost governance, MCP"],
])

# Appendix
h2("Appendix — References")
bullets([
    "Claude Managed Agents: platform.claude.com/docs/en/managed-agents/overview",
    "Claude Agent SDK: code.claude.com/docs/en/agent-sdk/overview",
    "Anthropic release notes (May 2026): releasebot.io/updates/anthropic",
    "Repo: github.com/hkuangedu/apps-Archway",
    "Finnhub API: finnhub.io/docs/api",
])

doc.save(OUT)
print(f"Wrote {OUT}")
